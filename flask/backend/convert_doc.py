import os
import re
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import subprocess
import sys

### --- FOR PDF --- ###
# def word_to_pdf(input_docx, output_pdf=None):
#     try:
#         if output_pdf is None:
#             output_pdf = os.path.splitext(input_docx)[0] + ".pdf"

#         output_dir = os.path.dirname(output_pdf)
#         os.makedirs(output_dir, exist_ok=True)

#         cmd = [
#             "soffice",
#             "--headless",
#             "--convert-to", "pdf",
#             "--outdir", output_dir,
#             input_docx
#         ]

#         result = subprocess.run(
#             cmd,
#             stdout=subprocess.PIPE,
#             stderr=subprocess.PIPE,
#             text=True,
#             encoding="utf-8",
#             errors="replace"
#         )

#         generated_pdf = os.path.join(
#             output_dir,
#             os.path.splitext(os.path.basename(input_docx))[0] + ".pdf"
#         )

#         if result.returncode != 0:
#             raise RuntimeError(
#                 f"LibreOffice PDF conversion failed.\nSTDOUT:\n{result.stdout}\nSTDERR:\n{result.stderr}"
#             )

#         if not os.path.exists(generated_pdf):
#             raise RuntimeError(
#                 f"LibreOffice finished but PDF was not created.\nSTDOUT:\n{result.stdout}\nSTDERR:\n{result.stderr}"
#             )

#         if os.path.abspath(generated_pdf) != os.path.abspath(output_pdf):
#             if os.path.exists(output_pdf):
#                 os.remove(output_pdf)
#             os.replace(generated_pdf, output_pdf)

#         return output_pdf

#     except Exception as e:
#         print(f"PDF conversion skipped due to error: {e}")



def word_to_pdf(input_docx, output_pdf=None):
    try:
        cmd = [
            sys.executable,
            "-c",
            f"from docx2pdf import convert; convert(r'{input_docx}', r'{output_pdf}')"
        ]

        result = subprocess.run(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE
        )

        # Ignore strange Windows exit codes but log them
        if result.returncode != 0:
            print("PDF conversion returned non-zero exit code but continuing...")
            print(result.stderr.decode(errors="ignore"))

    except Exception as e:
        print(f"PDF conversion skipped due to error: {e}")

### --- FOR WORD --- ###
ANCHOR_RE = re.compile(r"^\s*<a\s+id=['\"].*?['\"]\s*>\s*</a>\s*$", re.IGNORECASE)
BR_RE = re.compile(r"^\s*<br\s*/?>\s*$", re.IGNORECASE)

MD_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
INLINE_CODE_RE = re.compile(r"`([^`]+)`")

FILENAME_RE = re.compile(
    r"^[A-Za-z0-9_.\\/-]+\.(aspx|vb|cs|sql|config|json|md|txt)$",
    re.IGNORECASE
)

def tighten_styles(doc: Document):
    def tighten(style_name: str, before=0, after=0, line=1.0):
        try:
            s = doc.styles[style_name]
            pf = s.paragraph_format
            pf.space_before = Pt(before)
            pf.space_after = Pt(after)
            pf.line_spacing = line
        except KeyError:
            pass

    tighten("Normal", before=0, after=0, line=1.0)
    tighten("Heading 1", before=10, after=2, line=1.0)
    tighten("Heading 2", before=6, after=2, line=1.0)
    tighten("Heading 3", before=4, after=2, line=1.0)
    tighten("List Bullet", before=0, after=0, line=1.0)

def strip_md_heading_prefix(s: str) -> str:
    return re.sub(r"^\s*#{1,6}\s+", "", s).strip()

def clean_inline(text: str) -> str:
    text = MD_LINK_RE.sub(r"\1", text)
    text = INLINE_CODE_RE.sub(r"\1", text)
    return text

def add_runs_with_bold(paragraph, text: str):
    text = clean_inline(text)
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        r = paragraph.add_run(m.group(1))
        r.bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])

def looks_like_chapter_title(line: str) -> bool:
    ll = line.lower().strip()
    if ll.startswith("page: "):
        return True
    if ll.startswith("procedure: "):
        return True
    if FILENAME_RE.match(line.strip()):
        return True
    return False

def is_md_h1(line: str) -> bool:
    return line.strip().startswith("# ")

def is_md_h2(line: str) -> bool:
    return line.strip().startswith("## ")

def is_md_h3(line: str) -> bool:
    return line.strip().startswith("### ")

def md_heading_text(line: str) -> str:
    return clean_inline(line.strip().lstrip("#").strip())

def is_document_title(line: str) -> bool:
    return is_md_h1(line) and md_heading_text(line).lower().endswith("technical documentation")

def md_bullet_text(line: str) -> str:
    s = line.strip()

    if s.startswith("•"):
        return clean_inline(s.lstrip("•").strip())

    if s.startswith("- "):
        return clean_inline(s[2:].strip())

    if s.startswith("* "):
        return clean_inline(s[2:].strip())

    if s.startswith("+ "):
        return clean_inline(s[2:].strip())

    return ""

def find_toc_range(lines):
    start = None
    end = None

    for i, raw in enumerate(lines):
        line = raw.strip()

        if ANCHOR_RE.match(line):
            continue

        if start is None and is_md_h2(line) and md_heading_text(line).lower() == "table of contents":
            start = i
            continue

        if start is not None:
            # End TOC at the explicit separator before body content
            if BR_RE.match(line):
                end = i + 1
                break

    if start is None:
        return None, None
    if end is None:
        end = len(lines)
    return start, end

def collect_toc_structure(md_text: str):
    lines = md_text.splitlines()
    start, end = find_toc_range(lines)
    if start is None:
        return []

    in_toc = False
    current_cat = None
    order = []
    buckets = {}

    for i in range(start, end):
        raw = lines[i]
        line = raw.strip()

        if ANCHOR_RE.match(line) or BR_RE.match(line):
            continue

        if not in_toc:
            in_toc = True
            continue

        if is_md_h3(line):
            current_cat = md_heading_text(line)
            if current_cat not in buckets:
                buckets[current_cat] = []
                order.append(current_cat)
            continue

        bt = md_bullet_text(raw)
        if bt and current_cat:
            buckets[current_cat].append(bt)

    return [(cat, buckets.get(cat, [])) for cat in order]

def toc_item_to_chapter_title(item: str) -> str:
    return clean_inline(strip_md_heading_prefix(item)).strip()

def slug_bookmark(text: str) -> str:
    base = re.sub(r"[^A-Za-z0-9_]+", "_", text.strip())
    if not base:
        base = "Section"
    if not re.match(r"^[A-Za-z_]", base):
        base = "S_" + base
    return base[:40]

def add_bookmark(paragraph, bookmark_name: str, bookmark_id: int):
    p = paragraph._p
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), bookmark_name)

    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))

    p.insert(0, start)
    p.append(end)

def add_internal_hyperlink(paragraph, text: str, bookmark_name: str):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)
    hyperlink.set(qn("w:history"), "1")

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rPr.append(color)

    r.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    r.append(t)

    hyperlink.append(r)
    paragraph._p.append(hyperlink)

def is_table_line(s: str) -> bool:
    s = s.strip()
    return s.startswith("|") and s.endswith("|") and "|" in s[1:-1]

def is_table_sep(s: str) -> bool:
    s = s.strip()
    if not is_table_line(s):
        return False
    cells = [c.strip() for c in s.strip("|").split("|")]
    return all(re.fullmatch(r":?-{3,}:?", c or "") for c in cells)

def parse_table(lines, i):
    if i + 1 >= len(lines):
        return None, i

    header = lines[i].rstrip("\n")
    sep = lines[i + 1].rstrip("\n")
    if not (is_table_line(header) and is_table_sep(sep)):
        return None, i

    rows = []
    rows.append([c.strip() for c in header.strip().strip("|").split("|")])

    j = i + 2
    while j < len(lines):
        row = lines[j].rstrip("\n")
        if not is_table_line(row):
            break
        rows.append([c.strip() for c in row.strip().strip("|").split("|")])
        j += 1

    return rows, j

def markdown_to_word(md_text: str, output_path: str):
    doc = Document()
    tighten_styles(doc)

    lines = md_text.splitlines()
    toc = collect_toc_structure(md_text)

    bookmark_map = {}
    category_bookmark_map = {}
    chapter_to_category = {}
    used = set()

    for category, items in toc:
        base = slug_bookmark(category)
        name = base
        k = 1
        while name in used:
            k += 1
            name = f"{base}_{k}"
        used.add(name)
        category_bookmark_map[category] = name

        for item in items:
            chapter_title = toc_item_to_chapter_title(item)
            chapter_title = clean_inline(strip_md_heading_prefix(chapter_title)).strip()

            chapter_to_category[chapter_title] = category

            base = slug_bookmark(chapter_title)
            name = base
            k = 1
            while name in used:
                k += 1
                name = f"{base}_{k}"
            used.add(name)
            bookmark_map[chapter_title] = name

    doc_title = "Technical Documentation"
    if lines:
        first_nonempty = next((ln.strip() for ln in lines if ln.strip()), "")
        if first_nonempty and is_document_title(first_nonempty):
            doc_title = md_heading_text(first_nonempty)

    doc.add_heading(doc_title, level=1)

    doc.add_heading("Table of Contents", level=2)
    for category, items in toc:
        p_cat = doc.add_heading("", level=3)
        cat_bm = category_bookmark_map.get(category)
        if cat_bm:
            add_internal_hyperlink(p_cat, category, cat_bm)
        else:
            add_runs_with_bold(p_cat, category)

        for item in items:
            chapter_title = toc_item_to_chapter_title(item)
            chapter_title = clean_inline(strip_md_heading_prefix(chapter_title)).strip()

            bm = bookmark_map.get(chapter_title)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Pt(18)

            if bm:
                add_internal_hyperlink(p, clean_inline(item), bm)
            else:
                add_runs_with_bold(p, clean_inline(item))

    doc.add_page_break()

    toc_start, toc_end = find_toc_range(lines)

    prev_blank = True
    bookmark_id = 1
    have_written_any_chapter = False
    seen_categories_in_body = set()

    def add_single_blank():
        nonlocal prev_blank
        if not prev_blank:
            doc.add_paragraph("")
            prev_blank = True

    i = 0
    while i < len(lines):
        if toc_start is not None and toc_end is not None and toc_start <= i < toc_end:
            i = toc_end
            continue

        raw = lines[i].rstrip("\n")
        line = raw.strip()

        content_line = strip_md_heading_prefix(line)
        content_line_clean = clean_inline(content_line).strip()

        if ANCHOR_RE.match(line):
            i += 1
            continue

        if BR_RE.match(line):
            add_single_blank()
            i += 1
            continue

        if line == "":
            add_single_blank()
            i += 1
            continue

        if is_document_title(line):
            i += 1
            continue

        if is_md_h1(line) and not is_document_title(line):
            p = doc.add_heading("", level=1)
            add_runs_with_bold(p, content_line_clean)

            bm = bookmark_map.get(content_line_clean)
            if bm:
                add_bookmark(p, bm, bookmark_id)
                bookmark_id += 1

            prev_blank = False
            i += 1
            continue

        if is_md_h3(line) and not looks_like_chapter_title(content_line_clean):
            p = doc.add_heading("", level=3)
            add_runs_with_bold(p, content_line_clean)
            prev_blank = False
            i += 1
            continue

        if is_md_h2(line) and md_heading_text(line).lower() != "table of contents" and not looks_like_chapter_title(content_line_clean):
            p = doc.add_heading("", level=2)
            add_runs_with_bold(p, content_line_clean)

            cat_bm = category_bookmark_map.get(content_line_clean)
            if cat_bm:
                add_bookmark(p, cat_bm, bookmark_id)
                bookmark_id += 1

            prev_blank = False
            i += 1
            continue

        if looks_like_chapter_title(content_line_clean):
            title = content_line_clean

            if have_written_any_chapter:
                doc.add_paragraph("")
            have_written_any_chapter = True

            p = doc.add_heading("", level=1)
            add_runs_with_bold(p, title)

            bm = bookmark_map.get(title)
            if bm:
                add_bookmark(p, bm, bookmark_id)
                bookmark_id += 1

            prev_blank = False
            i += 1
            continue

        if re.match(r"^\d+\.\s+", content_line_clean):
            m = re.match(r"^(\d+)\.\s+(.*)", content_line_clean)
            if m:
                num = m.group(1)
                text = m.group(2)
                p = doc.add_paragraph()
                p.add_run(f"{num}. ")
                add_runs_with_bold(p, text)
            else:
                text = re.sub(r"^\d+\.\s+", "", content_line_clean)
                p = doc.add_paragraph()
                add_runs_with_bold(p, text)
            prev_blank = False
            i += 1
            continue

        table_rows, next_i = parse_table(lines, i)
        if table_rows is not None:
            cols = max(len(r) for r in table_rows)
            for r in table_rows:
                while len(r) < cols:
                    r.append("")

            t = doc.add_table(rows=len(table_rows), cols=cols)
            t.style = "Table Grid"

            for r_idx, row in enumerate(table_rows):
                for c_idx, cell_text in enumerate(row):
                    cell = t.cell(r_idx, c_idx)
                    p0 = cell.paragraphs[0]
                    for run in p0.runs:
                        run.text = ""
                    add_runs_with_bold(p0, cell_text)
                    if r_idx == 0:
                        for run in p0.runs:
                            run.bold = True

            prev_blank = False
            i = next_i
            continue

        bt = md_bullet_text(raw)
        if bt:
            indent_level = 0
            if raw.startswith("    ") or raw.startswith("\t"):
                indent_level = 1

            style = "List Bullet" if indent_level == 0 else "List Bullet 2"

            p = doc.add_paragraph(style=style)
            add_runs_with_bold(p, bt)

            prev_blank = False
            i += 1
            continue

        if re.fullmatch(r"-{3,}", content_line_clean):
            add_single_blank()
            i += 1
            continue

        p = doc.add_paragraph()
        add_runs_with_bold(p, clean_inline(strip_md_heading_prefix(raw)))
        prev_blank = False
        i += 1

    doc.save(output_path)

def main():
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    final_dir = os.path.join(BASE_DIR, "final_output")

    md_path = os.path.join(final_dir, "Technical_Documentation.md")
    docx_path = os.path.join(final_dir, "Technical_Documentation.docx")
    pdf_path = os.path.join(final_dir, "Technical_Documentation.pdf")

    if not os.path.exists(md_path):
        raise FileNotFoundError(f"Markdown file not found: {md_path}")

    os.makedirs(final_dir, exist_ok=True)

    with open(md_path, "r", encoding="utf-8") as f:
        md_text = f.read()

    markdown_to_word(md_text, docx_path)
    print(f"Word document created: {docx_path}")

    word_to_pdf(docx_path, pdf_path)

    if os.path.exists(pdf_path):
        print(f"PDF created: {pdf_path}")
    else:
        print("PDF was not created.")

if __name__ == "__main__":
    main()